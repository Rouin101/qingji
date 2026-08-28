"""Write trusted project exports as Markdown, DOCX, or PDF."""

from __future__ import annotations

import html
import re
from collections.abc import Iterable
from pathlib import Path
from typing import Any

from .export import export_project_markdown


EXPORT_FORMATS = ("markdown", "docx", "pdf")
EXPORT_FORMAT_LABELS = {
    "markdown": "Markdown (.md)",
    "docx": "Word 文档 (.docx)",
    "pdf": "PDF (.pdf)",
}


def project_output_root() -> Path:
    """Return the repository-local directory for user-generated exports."""

    return Path(__file__).resolve().parents[1] / "output"


def _normalise_formats(formats: Iterable[str]) -> tuple[str, ...]:
    selected = tuple(dict.fromkeys(str(item).lower() for item in formats))
    invalid = sorted(set(selected).difference(EXPORT_FORMATS))
    if invalid:
        raise ValueError(f"不支持的导出格式：{'、'.join(invalid)}")
    if not selected:
        raise ValueError("请至少选择一种导出格式。")
    return selected


def _export_path(
    project_id: int, export_format: str, output_root: Path | None
) -> Path:
    suffix = {"markdown": ".md", "docx": ".docx", "pdf": ".pdf"}[export_format]
    root = Path(output_root) if output_root is not None else project_output_root()
    directory = root / export_format
    directory.mkdir(parents=True, exist_ok=True)
    return directory / f"青迹_项目{int(project_id)}_可信导出{suffix}"


def _atomic_write_text(destination: Path, content: str) -> None:
    temporary = destination.with_suffix(destination.suffix + ".tmp")
    temporary.write_text(content, encoding="utf-8", newline="\n")
    temporary.replace(destination)


def _set_word_font(style: Any, *, name: str, size: float, color: str, bold: bool) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(key), name)


def _word_style(
    document: Any,
    name: str,
    *,
    size: float,
    color: str,
    bold: bool,
    before: float,
    after: float,
    line_spacing: float = 1.1,
    left_indent: float | None = None,
    first_line_indent: float | None = None,
) -> Any:
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Inches, Pt, RGBColor

    try:
        style = document.styles[name]
    except KeyError:
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    _set_word_font(
        style,
        name="Microsoft YaHei",
        size=size,
        color=color,
        bold=bold,
    )
    paragraph = style.paragraph_format
    paragraph.space_before = Pt(before)
    paragraph.space_after = Pt(after)
    paragraph.line_spacing = line_spacing
    if left_indent is not None:
        paragraph.left_indent = Inches(left_indent)
    if first_line_indent is not None:
        paragraph.first_line_indent = Inches(first_line_indent)
    return style


def _append_markdown_runs(paragraph: Any, text: str) -> None:
    """Copy the small bold subset used by the trusted Markdown renderer."""

    pieces = re.split(r"(\*\*.+?\*\*)", text)
    for piece in pieces:
        if not piece:
            continue
        bold = piece.startswith("**") and piece.endswith("**")
        run = paragraph.add_run(piece[2:-2] if bold else piece)
        run.bold = bold


def render_markdown_to_docx(markdown: str, destination: str | Path, *, visual_path: str | Path | None = None) -> Path:
    """Render trusted Markdown using the standard_business_brief token set."""

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    target = Path(destination)
    target.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    # standard_business_brief, with a Microsoft YaHei override for Chinese text.
    normal = _word_style(
        document, "Normal", size=11, color="1F2937", bold=False,
        before=0, after=6, line_spacing=1.1,
    )
    title = _word_style(
        document, "Qingji Title", size=24, color="0B2545", bold=True,
        before=0, after=4, line_spacing=1.1,
    )
    subtitle = _word_style(
        document, "Qingji Subtitle", size=11, color="667085", bold=False,
        before=0, after=18, line_spacing=1.1,
    )
    heading_1 = _word_style(
        document, "Qingji Heading 1", size=16, color="2E74B5", bold=True,
        before=16, after=8, line_spacing=1.1,
    )
    heading_2 = _word_style(
        document, "Qingji Heading 2", size=13, color="2E74B5", bold=True,
        before=12, after=6, line_spacing=1.1,
    )
    heading_3 = _word_style(
        document, "Qingji Heading 3", size=12, color="1F4D78", bold=True,
        before=8, after=4, line_spacing=1.1,
    )
    quote = _word_style(
        document, "Qingji Quote", size=10.5, color="475467", bold=False,
        before=2, after=6, line_spacing=1.1, left_indent=0.25,
    )
    _word_style(
        document, "Qingji Note", size=9.5, color="667085", bold=False,
        before=4, after=4, line_spacing=1.1,
    )

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    header_run = header.add_run("青迹｜可信证据导出")
    header_run.font.name = "Microsoft YaHei"
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = RGBColor(102, 112, 133)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer_run = footer.add_run("由青迹本地可信证据链流程生成")
    footer_run.font.name = "Microsoft YaHei"
    footer_run.font.size = Pt(8.5)
    footer_run.font.color.rgb = RGBColor(102, 112, 133)

    lines = markdown.splitlines()
    title_text = "青迹可信证据导出"
    for line in lines:
        match = re.match(r"^#\s+(.+)$", line)
        if match:
            title_text = match.group(1).strip()
            break
    document.core_properties.title = title_text
    document.core_properties.subject = "当前项目的可信证据、核验结论与补证缺口"
    document.core_properties.author = "青迹"

    title_paragraph = document.add_paragraph(style=title)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _append_markdown_runs(title_paragraph, title_text)
    subtitle_paragraph = document.add_paragraph(style=subtitle)
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_paragraph.add_run("项目成果、证据与缺口可信导出")
    if visual_path and Path(visual_path).is_file():
        document.add_picture(str(visual_path), width=Inches(6.3))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        visual_caption = document.add_paragraph("项目概览图与材料时间线", style="Qingji Note")
        visual_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    skipped_title = False
    for raw_line in lines:
        line = raw_line.rstrip()
        if not line or line == "---":
            continue
        match = re.match(r"^(#{1,3})\s+(.+)$", line)
        if match:
            level = len(match.group(1))
            text = match.group(2).strip()
            if level == 1 and not skipped_title:
                skipped_title = True
                continue
            style = {1: heading_1, 2: heading_2, 3: heading_3}[level]
            paragraph = document.add_paragraph(style=style)
            paragraph.paragraph_format.keep_with_next = True
            _append_markdown_runs(paragraph, text)
            continue
        if line.startswith("> "):
            paragraph = document.add_paragraph(style=quote)
            _append_markdown_runs(paragraph, line[2:].strip())
            continue
        bullet = re.match(r"^\s*-\s+(.+)$", line)
        if bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.line_spacing = 1.167
            _append_markdown_runs(paragraph, bullet.group(1).strip())
            continue
        paragraph = document.add_paragraph(style=normal)
        _append_markdown_runs(paragraph, line)

    temporary = target.with_suffix(target.suffix + ".tmp")
    document.save(temporary)
    temporary.replace(target)
    return target


_PDF_FONT_NAME = "QingjiChinese"


def _register_pdf_font() -> str:
    """Embed a Windows Chinese TrueType font when it is available."""

    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase.ttfonts import TTFont

    try:
        pdfmetrics.getFont(_PDF_FONT_NAME)
    except KeyError:
        font_path = Path("C:/Windows/Fonts/msyh.ttc")
        if font_path.is_file():
            pdfmetrics.registerFont(TTFont(_PDF_FONT_NAME, str(font_path), subfontIndex=0))
        else:
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
            return "STSong-Light"
    return _PDF_FONT_NAME


def _pdf_styles() -> dict[str, Any]:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet

    font_name = _register_pdf_font()
    base = getSampleStyleSheet()["BodyText"]
    styles = {
        "title": ParagraphStyle(
            "QingjiPdfTitle", parent=base, fontName=font_name, fontSize=22,
            leading=29, textColor=colors.HexColor("#0B2545"), alignment=TA_CENTER,
            spaceAfter=6,
        ),
        "subtitle": ParagraphStyle(
            "QingjiPdfSubtitle", parent=base, fontName=font_name, fontSize=10.5,
            leading=15, textColor=colors.HexColor("#667085"), alignment=TA_CENTER,
            spaceAfter=18,
        ),
        "body": ParagraphStyle(
            "QingjiPdfBody", parent=base, fontName=font_name, fontSize=10.5,
            leading=16, textColor=colors.HexColor("#1F2937"), spaceAfter=7,
        ),
        "h1": ParagraphStyle(
            "QingjiPdfH1", parent=base, fontName=font_name, fontSize=15,
            leading=22, textColor=colors.HexColor("#2E74B5"), spaceBefore=14,
            spaceAfter=7,
        ),
        "h2": ParagraphStyle(
            "QingjiPdfH2", parent=base, fontName=font_name, fontSize=12.5,
            leading=19, textColor=colors.HexColor("#2E74B5"), spaceBefore=11,
            spaceAfter=5,
        ),
        "h3": ParagraphStyle(
            "QingjiPdfH3", parent=base, fontName=font_name, fontSize=11.5,
            leading=17, textColor=colors.HexColor("#1F4D78"), spaceBefore=7,
            spaceAfter=4,
        ),
        "quote": ParagraphStyle(
            "QingjiPdfQuote", parent=base, fontName=font_name, fontSize=10,
            leading=15, textColor=colors.HexColor("#475467"), leftIndent=18,
            borderColor=colors.HexColor("#D0D5DD"), borderWidth=1,
            borderPadding=7, borderLeft=True, spaceAfter=7,
        ),
    }
    return styles


def _pdf_header_footer(canvas: Any, document: Any) -> None:
    from reportlab.lib import colors

    canvas.saveState()
    canvas.setStrokeColor(colors.HexColor("#D0D5DD"))
    canvas.line(document.leftMargin, 758, 612 - document.rightMargin, 758)
    canvas.setFillColor(colors.HexColor("#667085"))
    canvas.setFont(_register_pdf_font(), 8.5)
    canvas.drawRightString(540, 765, "青迹｜可信证据导出")
    canvas.drawCentredString(306, 34, f"第 {document.page} 页｜由青迹本地可信证据链流程生成")
    canvas.restoreState()


def render_markdown_to_pdf(markdown: str, destination: str | Path, *, visual_path: str | Path | None = None) -> Path:
    """Render trusted Markdown as a readable, local PDF."""

    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.platypus import Image as PdfImage, ListFlowable, ListItem, Paragraph, SimpleDocTemplate

    target = Path(destination)
    target.parent.mkdir(parents=True, exist_ok=True)
    temporary = target.with_suffix(target.suffix + ".tmp")
    styles = _pdf_styles()
    document = SimpleDocTemplate(
        str(temporary), pagesize=letter,
        leftMargin=inch, rightMargin=inch, topMargin=0.78 * inch, bottomMargin=0.72 * inch,
        title="青迹可信证据导出", author="青迹",
    )
    story: list[Any] = []
    if visual_path and Path(visual_path).is_file():
        visual = PdfImage(str(visual_path))
        visual.drawWidth = 6.25 * inch
        visual.drawHeight = visual.imageHeight * visual.drawWidth / visual.imageWidth
        story.append(visual)
        story.append(Paragraph("项目概览图与材料时间线", styles["subtitle"]))
    lines = markdown.splitlines()
    skipped_title = False
    for raw_line in lines:
        line = raw_line.rstrip()
        if not line or line == "---":
            continue
        match = re.match(r"^(#{1,3})\s+(.+)$", line)
        if match:
            level = len(match.group(1))
            value = html.escape(match.group(2).strip())
            if level == 1 and not skipped_title:
                story.append(Paragraph(value, styles["title"]))
                story.append(Paragraph("项目成果、证据与缺口可信导出", styles["subtitle"]))
                skipped_title = True
            else:
                story.append(Paragraph(value, styles[f"h{level}"]))
            continue
        if line.startswith("> "):
            story.append(Paragraph(html.escape(line[2:].strip()), styles["quote"]))
            continue
        bullet = re.match(r"^\s*-\s+(.+)$", line)
        if bullet:
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(html.escape(bullet.group(1).strip()), styles["body"]))],
                    bulletType="bullet", leftIndent=18, bulletFontName="Helvetica",
                    bulletFontSize=8, spaceAfter=2,
                )
            )
            continue
        story.append(Paragraph(html.escape(line), styles["body"]))
    if not skipped_title:
        story.insert(0, Paragraph("青迹可信证据导出", styles["title"]))
        story.insert(1, Paragraph("项目成果、证据与缺口可信导出", styles["subtitle"]))
    document.build(story, onFirstPage=_pdf_header_footer, onLaterPages=_pdf_header_footer)
    temporary.replace(target)
    return target


def write_export_files(
    markdown: str,
    project_id: int,
    formats: Iterable[str] = EXPORT_FORMATS,
    *,
    output_root: str | Path | None = None,
) -> dict[str, Path]:
    """Write requested formats below ``output/<format>/`` and return paths."""

    selected = _normalise_formats(formats)
    root = Path(output_root) if output_root is not None else None
    written: dict[str, Path] = {}
    for export_format in selected:
        destination = _export_path(project_id, export_format, root)
        if export_format == "markdown":
            _atomic_write_text(destination, markdown)
        elif export_format == "docx":
            render_markdown_to_docx(markdown, destination)
        else:
            render_markdown_to_pdf(markdown, destination)
        written[export_format] = destination
    return written


def export_project_files(
    db: Any,
    project_id: int,
    formats: Iterable[str] = EXPORT_FORMATS,
    *,
    output_root: str | Path | None = None,
) -> dict[str, Path]:
    """Build the trusted source once, then write the selected file formats."""

    return write_export_files(
        export_project_markdown(db, project_id), project_id, formats,
        output_root=output_root,
    )